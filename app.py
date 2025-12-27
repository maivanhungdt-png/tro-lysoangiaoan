
import streamlit as st
import google.generativeai as genai
from PIL import Image
import tempfile
import os
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm

# =========================================================
# CHUYá»‚N CÃ”NG THá»¨C SANG CHUáº¨N MASSIVEMARK (BIBCIT)
# =========================================================
def convert_math_for_massivemark(text: str) -> str:
    text = re.sub(
        r'\[MATH\](.*?)\[/MATH\]',
        lambda m: r'\(' + m.group(1).strip() + r'\)',
        text,
        flags=re.DOTALL
    )
    pattern = r'(?<!\\\()(\b(?:\\frac\{.*?\}\{.*?\}|\\sqrt\{.*?\}|[0-9a-zA-Z]+(?:\^[0-9a-zA-Z]+)?\s*(?:=|>|<|â‰¥|â‰¤)\s*[0-9a-zA-Z]+(?:\^[0-9a-zA-Z]+)?))'
    text = re.sub(pattern, r'\\(\1\\)', text)
    return text

# =========================================================
# Xá»¬ LÃ CÃ”NG THá»¨C TOÃN THCS â€“ CHUáº¨N SGK + MATHTYPE
# =========================================================
def auto_wrap_math(text: str) -> str:
    pattern = r'(?<!\[MATH\])(\b(?:\\frac\{.*?\}\{.*?\}|\\sqrt\{.*?\}|[0-9a-zA-Z]+(?:\^[0-9a-zA-Z]+)?\s*(?:=|>|<|â‰¥|â‰¤)\s*[0-9a-zA-Z]+(?:\^[0-9a-zA-Z]+)?))'
    return re.sub(pattern, r'[MATH]\1[/MATH]', text)

def process_math_blocks(text: str) -> str:
    def repl(match):
        expr = match.group(1).strip()
        expr = re.sub(r'\$(.*?)\$', r'\1', expr)
        return expr
    return re.sub(r'\[MATH\](.*?)\[/MATH\]', repl, text, flags=re.DOTALL)

# --- WORD HELPERS ---
def add_formatted_text(paragraph, text):
    # FIX: do not access paragraph.style.font directly
    paragraph.style = paragraph.part.document.styles['Normal']
    parts = re.split(r'(\*\*.*?\*\*)', text) 
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean = part[2:-2]
            run = paragraph.add_run(clean)
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def create_doc_stable(content, ten_bai, lop):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.2

    head = doc.add_heading(f'Káº¾ HOáº CH BÃ€I Dáº Y: {ten_bai.upper()}', 0)
    head.alignment = 1
    for run in head.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    p_lop = doc.add_paragraph(f'Lá»›p: {lop}')
    p_lop.alignment = 1
    p_lop.runs[0].bold = True
    doc.add_paragraph("-" * 60).alignment = 1

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('#'):
            line = line.replace('#', '').strip()

        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            valid_rows = [r for r in table_lines if '---' not in r]
            if len(valid_rows) >= 2:
                cols = len(valid_rows[0].split('|')) - 2
                table = doc.add_table(rows=len(valid_rows), cols=cols)
                table.style = 'Table Grid'
                for r_idx, r_text in enumerate(valid_rows):
                    cells_data = r_text.split('|')[1:-1]
                    for c_idx, cell_text in enumerate(cells_data):
                        cell = table.cell(r_idx, c_idx)
                        cell._element.clear_content()
                        raw = cell_text.strip().replace('<br>', '\n').replace('<br/>', '\n')
                        for sub in raw.split('\n'):
                            if not sub.strip():
                                continue
                            p = cell.add_paragraph()
                            if r_idx == 0:
                                p.alignment = 1
                                r = p.add_run(sub.replace('**',''))
                                r.bold = True
                            else:
                                add_formatted_text(p, sub.strip())
            continue

        if not line:
            i += 1
            continue

        if re.match(r'^(I\.|II\.|III\.|IV\.|V\.)', line) or (re.match(r'^\d+\.', line) and len(line) < 50):
            p = doc.add_paragraph(line.replace('**',''))
            p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(14)
        elif line.startswith('- '):
            p = doc.add_paragraph("â€“ ")
            add_formatted_text(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        i += 1

    return doc

# --- STREAMLIT UI ---
st.set_page_config(page_title="Trá»£ lÃ½ GiÃ¡o Ã¡n NLS", page_icon="ğŸ“˜", layout="centered")

# ================== UI HEADER (GIá»NG Báº¢N CÅ¨) ==================
st.markdown("""
<style>
    [data-testid="stAppViewContainer"] { background-color: #f4f6f9; }

    .main-header {
        background: linear-gradient(135deg, #004e92 0%, #000428 100%);
        padding: 30px;
        border-radius: 15px;
        text-align: center;
        color: white !important;
        margin-bottom: 30px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .main-header h1 {
        color: white !important;
        margin: 0;
        font-family: 'Segoe UI', sans-serif;
        font-size: 2rem;
    }
    .main-header p {
        color: #e0e0e0 !important;
        margin-top: 10px;
        font-style: italic;
    }
    .section-header {
        color: #004e92;
        border-bottom: 2px solid #ddd;
        padding-bottom: 5px;
        margin-top: 20px;
        margin-bottom: 15px;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>ğŸ“˜ TRá»¢ LÃ SOáº N GIÃO ÃN KHUNG NÄ‚NG Lá»°C Sá» Tá»° Äá»˜NG (NLS)</h1>
    <p>TÃ¡c giáº£: Mai VÄƒn HÃ¹ng - TrÆ°á»ng THCS Äá»“ng YÃªn - SÄT: 0941037116</p>
</div>
""", unsafe_allow_html=True)
# ===============================================================

FILE_KHUNG_NANG_LUC = "khungnanglucso.pdf"

st.markdown("<h1>ğŸ“˜ TRá»¢ LÃ SOáº N GIÃO ÃN</h1>", unsafe_allow_html=True)

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    api_key = st.text_input("Nháº­p API Key:", type="password")

if api_key:
    genai.configure(api_key=api_key)


# ================== KHUNG NÄ‚NG Lá»°C Sá» ==================
st.markdown('<div class="section-header">ğŸ“‚ 1. TÃ€I LIá»†U NGUá»’N</div>', unsafe_allow_html=True)

has_framework = False
if os.path.exists(FILE_KHUNG_NANG_LUC):
    st.success(f"âœ… ÄÃ£ tÃ­ch há»£p khung nÄƒng lá»±c sá»‘: {FILE_KHUNG_NANG_LUC}")
    has_framework = True
else:
    st.info("â„¹ï¸ ChÆ°a cÃ³ file khung nÄƒng lá»±c sá»‘ (khungnanglucso.pdf).")

uploaded_files = st.file_uploader(
    "Táº£i áº¢nh/PDF bÃ i dáº¡y (kÃ©o tháº£ vÃ o Ä‘Ã¢y):",
    type=["jpg", "png", "pdf"],
    accept_multiple_files=True
)

if uploaded_files:
    st.caption("ğŸ‘ï¸ Xem trÆ°á»›c tÃ i liá»‡u:")
    cols = st.columns(3)
    for i, f in enumerate(uploaded_files):
        if f.type in ["image/jpeg", "image/png"]:
            with cols[i % 3]:
                st.image(f, caption=f.name)
        else:
            with cols[i % 3]:
                st.info(f"ğŸ“„ {f.name}")
# ======================================================

uploaded_files = st.file_uploader("Táº£i áº¢nh/PDF bÃ i dáº¡y:", type=["jpg","png","pdf"], accept_multiple_files=True)

output_mode = st.radio("ğŸ§® Chá»n cÃ¡ch xá»­ lÃ½ cÃ´ng thá»©c:", ["Word / MathType", "Copy MassiveMark (BibCit)"], index=1)
lop = st.text_input("ğŸ“š Lá»›p:", "Lá»›p 6")
ten_bai = st.text_input("ğŸ“Œ TÃªn bÃ i há»c:", "")
noidung_bosung = st.text_area("âœï¸ Ghi chÃº thÃªm:", height=100)

if st.button("ğŸš€ SOáº N GIÃO ÃN NGAY"):
    if not api_key:
        st.error("Thiáº¿u API Key")
    else:
        temp_paths = []
        try:
            with st.spinner("AI Ä‘ang soáº¡n giÃ¡o Ã¡n..."):
                model = genai.GenerativeModel('gemini-2.5-flash-lite-preview-09-2025')
                prompt = f"""
ÄÃ³ng vai lÃ  má»™t GiÃ¡o viÃªn THCS vá»›i hÆ¡n 15 nÄƒm kinh nghiá»‡m dáº¡y há»c, am hiá»ƒu chÆ°Æ¡ng trÃ¬nh GDPT 2018.
Nhiá»‡m vá»¥: Soáº¡n Káº¿ hoáº¡ch bÃ i dáº¡y (GiÃ¡o Ã¡n) theo CÃ´ng vÄƒn 5512 cho bÃ i: "{ten_bai}" â€“ {lop}.

Dá»® LIá»†U Äáº¦U VÃ€O:
- CÃ¡c hÃ¬nh áº£nh/PDF SGK vÃ  tÃ i liá»‡u Ä‘Ã­nh kÃ¨m (náº¿u cÃ³): dÃ¹ng Ä‘á»ƒ trÃ­ch xuáº¥t CHÃNH XÃC kiáº¿n thá»©c.
- Ghi chÃº bá»• sung cá»§a giÃ¡o viÃªn: "{noidung_bosung}".

YÃŠU Cáº¦U Báº®T BUá»˜C Vá»€ Cáº¤U TRÃšC (CÃ”NG VÄ‚N 5512):
I. Má»¥c tiÃªu
1. Vá» kiáº¿n thá»©c
2. Vá» nÄƒng lá»±c
   a) NÄƒng lá»±c Ä‘áº·c thÃ¹
   b) NÄƒng lá»±c chung
   c) TÃ­ch há»£p nÄƒng lá»±c sá»‘
3. Vá» pháº©m cháº¥t

II. Thiáº¿t bá»‹ dáº¡y há»c vÃ  há»c liá»‡u
1. GiÃ¡o viÃªn
2. Há»c sinh

III. Tiáº¿n trÃ¬nh dáº¡y há»c
Gá»“m ÄÃšNG 4 hoáº¡t Ä‘á»™ng:
- Hoáº¡t Ä‘á»™ng 1: Khá»Ÿi Ä‘á»™ng
- Hoáº¡t Ä‘á»™ng 2: HÃ¬nh thÃ nh kiáº¿n thá»©c má»›i
  + Hoáº¡t Ä‘á»™ng 2.1 (á»©ng vá»›i má»¥c 1 SGK)
  + Hoáº¡t Ä‘á»™ng 2.2 (á»©ng vá»›i má»¥c 2 SGK) (cÃ³ thá»ƒ thÃªm 2.3 náº¿u SGK cÃ³)
- Hoáº¡t Ä‘á»™ng 3: Luyá»‡n táº­p
- Hoáº¡t Ä‘á»™ng 4: Váº­n dá»¥ng

Vá»šI Má»–I HOáº T Äá»˜NG, TRÃŒNH BÃ€Y THEO THá»¨ Tá»°:
a) Má»¥c tiÃªu
b) Ná»™i dung
c) Sáº£n pháº©m
d) Tá»• chá»©c thá»±c hiá»‡n (chá»‰ ghi dÃ²ng nÃ y, sau Ä‘Ã³ Ä‘áº¿n báº£ng)

SAU Má»¤C d) Báº®T BUá»˜C CÃ“ 01 Báº¢NG 2 Cá»˜T:

| Hoáº¡t Ä‘á»™ng cá»§a giÃ¡o viÃªn vÃ  há»c sinh | Ghi báº£ng |
|---|---|
| â€¦ | â€¦ |

QUY Äá»ŠNH Báº¢NG (KHÃ”NG NGOáº I Lá»†):
- Má»—i hoáº¡t Ä‘á»™ng chá»‰ cÃ³ 01 báº£ng
- Má»—i báº£ng chá»‰ cÃ³ 02 hÃ ng
- Ná»™i dung trong Ã´ gá»™p báº±ng <br>
- KhÃ´ng dÃ¹ng gáº¡ch Ä‘áº§u dÃ²ng tá»± Ä‘á»™ng trong báº£ng

Cá»˜T â€œHoáº¡t Ä‘á»™ng cá»§a giÃ¡o viÃªn vÃ  há»c sinhâ€:
Pháº£i mÃ´ táº£ Äáº¦Y Äá»¦ 4 BÆ¯á»šC:
BÆ°á»›c 1: Chuyá»ƒn giao nhiá»‡m vá»¥
BÆ°á»›c 2: Thá»±c hiá»‡n nhiá»‡m vá»¥
BÆ°á»›c 3: BÃ¡o cÃ¡o, tháº£o luáº­n
BÆ°á»›c 4: Káº¿t luáº­n, nháº­n Ä‘á»‹nh

Cá»˜T â€œGhi báº£ngâ€:
- Ghi TOÃ€N Bá»˜ káº¿t quáº£ kiáº¿n thá»©c Ä‘Ãºng SGK
- CÃ³ thá»ƒ gá»“m: khÃ¡i niá»‡m, Ä‘á»‹nh nghÄ©a, vÃ­ dá»¥, bÃ i táº­p, lá»i giáº£i chi tiáº¿t

QUY Æ¯á»šC CÃ”NG THá»¨C TOÃN:
- Má»ŒI cÃ´ng thá»©c toÃ¡n pháº£i Ä‘áº·t trong [MATH] ... [/MATH]
- Chá»‰ dÃ¹ng LaTeX cÆ¡ báº£n THCS: \\frac, \\sqrt, ^
- KHÃ”NG dÃ¹ng $, $$, \\text, \\mathbb, Unicode Â² Â³ âˆš
- KhÃ´ng Ä‘á»ƒ kÃ½ hiá»‡u =, <, >, â‰¥, â‰¤ trong vÄƒn báº£n thÆ°á»ng

IV. Äiá»u chá»‰nh sau tiáº¿t dáº¡y

LÆ¯U Ã:
- KhÃ´ng dÃ¹ng kÃ½ tá»± #
- KhÃ´ng mÃ´ táº£ káº¿t quáº£ sÆ° pháº¡m
- KhÃ´ng láº·p láº¡i cÃ¢u chá»¯ má»¥c tiÃªu
- KhÃ´ng bá» trá»‘ng cá»™t â€œGhi báº£ngâ€
- BÃ¡m sÃ¡t SGK vÃ  tÃ i liá»‡u Ä‘Ã­nh kÃ¨m
"""
                input_data = [prompt]

                if uploaded_files:
                    for f in uploaded_files:
                        if f.type == "application/pdf":
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(f.getvalue())
                                temp_paths.append(tmp.name)
                            input_data.append(genai.upload_file(tmp.name))
                        else:
                            input_data.append(Image.open(f))

                response = model.generate_content(input_data)
                if not response or not response.text or not response.text.strip():
                    st.error("AI khÃ´ng sinh Ä‘Æ°á»£c ná»™i dung.")
                    st.stop()

                ket_qua_text = response.text
                if output_mode == "Copy MassiveMark (BibCit)":
                    ket_qua_text = convert_math_for_massivemark(ket_qua_text)
                else:
                    ket_qua_text = auto_wrap_math(ket_qua_text)
                    ket_qua_text = process_math_blocks(ket_qua_text)

        except Exception as e:
            st.error(f"CÃ³ lá»—i: {e}")
            st.stop()
        finally:
            for p in temp_paths:
                if os.path.exists(p):
                    os.remove(p)


# ================== HIá»‚N THá»Š & XUáº¤T Káº¾T QUáº¢ ==================
st.markdown("### ğŸ“„ Káº¾T QUáº¢ BÃ€I SOáº N")
st.markdown(f'<div class="lesson-plan-paper">{ket_qua_text}</div>', unsafe_allow_html=True)

st.markdown("### ğŸ“‹ COPY Ná»˜I DUNG")
st.text_area("BÃ´i Ä‘en (Ctrl+A) â†’ Copy (Ctrl+C)", ket_qua_text, height=350)

if output_mode == "Word / MathType":
    safe_name = re.sub(r'[\\/:*?"<>|]', '', ten_bai) or "GiaoAn"
    doc = create_doc_stable(ket_qua_text, ten_bai, lop)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.download_button(
        label="â¬‡ï¸ Táº¢I FILE WORD CHUáº¨N A4",
        data=buf,
        file_name=f"GiaoAn_{safe_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )
# =============================================================
        st.markdown("### ğŸ“„ Káº¾T QUáº¢")
        st.text_area("Káº¿t quáº£", ket_qua_text, height=400)

        if output_mode == "Word / MathType":
            safe = re.sub(r'[\\/:*?"<>|]', '', ten_bai) or "GiaoAn"
            doc = create_doc_stable(ket_qua_text, ten_bai, lop)
            buf = io.BytesIO()
            doc.save(buf); buf.seek(0)
            st.download_button("â¬‡ï¸ Táº£i Word", buf, file_name=f"GiaoAn_{safe}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
